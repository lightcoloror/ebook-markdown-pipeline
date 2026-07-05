from __future__ import annotations

import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[1]
GENERATED_SIDECAR_NAMES = {"document-vlm-result.json", "table-candidates.json"}
TOOL_CACHE = Path(
    os.environ.get(
        "EBOOK_CONVERTER_TOOL_CACHE",
        Path.home() / ".cache" / "ebook-markdown-pipeline",
    )
)


def default_paddleocr_command() -> str:
    configured = os.environ.get("EBOOK_CONVERTER_PADDLEOCR_COMMAND", "").strip().strip('"')
    if configured:
        return configured
    return shutil.which("paddleocr") or "paddleocr"


def main() -> int:
    parser = argparse.ArgumentParser(description="Run PaddleOCR-VL doc_parser for one image and normalize output to Markdown.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--output-dir", type=Path, default=None)
    parser.add_argument("--device", default="cpu")
    parser.add_argument("--pipeline-version", default="v1.6", choices=["v1", "v1.5", "v1.6"])
    parser.add_argument("--timeout", type=float, default=0.0)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    image = args.input.resolve()
    output = args.output.resolve()
    work_dir = (args.output_dir or output.parent / "paddleocr_vl_raw").resolve()

    env = os.environ.copy()
    configure_local_cache(env)
    command = [
        default_paddleocr_command(),
        "doc_parser",
        "--input",
        str(image),
        "--save_path",
        str(work_dir),
        "--pipeline_version",
        args.pipeline_version,
        "--device",
        args.device,
        "--engine",
        "transformers",
        "--use_chart_recognition",
        "True",
        "--use_ocr_for_image_block",
        "True",
        "--format_block_content",
        "True",
        "--merge_layout_blocks",
        "True",
        "--max_new_tokens",
        "2048",
    ]
    if args.dry_run:
        print(subprocess.list2cmdline(command))
        print(f"document_vlm_result={work_dir / 'document-vlm-result.json'}")
        print(f"table_candidates={work_dir / 'table-candidates.json'}")
        return 0
    work_dir.mkdir(parents=True, exist_ok=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    if not shutil.which(command[0]) and not Path(command[0]).exists():
        raise FileNotFoundError(f"paddleocr executable not found: {command[0]}")
    completed = subprocess.run(
        command,
        env=env,
        text=True,
        encoding="utf-8",
        errors="replace",
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        timeout=args.timeout if args.timeout > 0 else None,
        check=False,
    )
    candidate = pick_text_artifact(work_dir)
    if candidate:
        write_normalized_markdown(candidate, output, work_dir)
    else:
        output.write_text(
            "# PaddleOCR-VL output unavailable\n\n"
            f"Return code: {completed.returncode}\n\n"
            "```text\n"
            f"{completed.stdout[-4000:]}\n"
            "```\n",
            encoding="utf-8",
            newline="\n",
        )
    markdown = output.read_text(encoding="utf-8", errors="replace")
    write_paddleocr_vl_sidecars(
        work_dir,
        image,
        output,
        markdown,
        command,
        status="review" if completed.returncode == 0 else "failed",
    )
    if completed.returncode != 0:
        print(completed.stdout[-4000:], file=sys.stderr)
        return completed.returncode
    print(str(output))
    return 0


def write_paddleocr_vl_sidecars(raw_dir: Path, source: Path, output: Path, markdown: str, command: list[str] | None, *, status: str) -> list[Path]:
    raw_dir.mkdir(parents=True, exist_ok=True)
    blocks = raw_document_blocks(raw_dir)
    if not blocks and markdown.strip():
        blocks = [
            {
                "type": "markdown_text",
                "text_preview": markdown.strip()[:1000],
                "text_char_count": len(markdown.strip()),
                "origin": "wrapper_markdown",
            }
        ]
    tables = table_candidates_from_raw(raw_dir)
    artifacts: list[dict[str, Any]] = [
        {"type": "markdown", "path": str(output), "label": "PaddleOCR-VL Markdown"},
        {"type": "raw_output_dir", "path": str(raw_dir), "label": "PaddleOCR-VL raw output"},
    ]
    document_payload: dict[str, Any] = {
        "schema_version": "document-vlm-result-v1",
        "backend": "paddleocr_vl",
        "status": status,
        "mode": "doc_parser",
        "input": str(source),
        "markdown": str(output),
        "pages": [{"page": 1, "source": str(source), "blocks": blocks, "tables": tables}],
        "artifacts": artifacts,
        "warnings": [],
        "promotion_use": "document-VLM review side evidence; explicit PaddleOCR-VL experiments only",
    }
    if command:
        document_payload["command"] = command
    document_path = raw_dir / "document-vlm-result.json"
    document_path.write_text(json.dumps(document_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8", newline="\n")

    table_payload = {
        "schema_version": "table-candidates-v1",
        "backend": "paddleocr_vl",
        "status": status,
        "mode": "doc_parser",
        "input": str(source),
        "markdown": str(output),
        "pages": [{"page": 1, "source": str(source), "tables": tables}],
        "artifacts": artifacts,
        "warnings": [] if tables else ["no table candidates extracted from PaddleOCR-VL raw outputs"],
    }
    table_path = raw_dir / "table-candidates.json"
    table_path.write_text(json.dumps(table_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8", newline="\n")
    return [document_path, table_path]


def raw_document_blocks(root: Path) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for path in sorted(root.rglob("*.json")):
        if path.name in GENERATED_SIDECAR_NAMES:
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        except Exception:
            continue
        blocks.extend(document_blocks_from_json_value(data, path))
    return blocks


def document_blocks_from_json_value(value: Any, source_artifact: Path) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    if isinstance(value, dict):
        label = str(value.get("block_label") or value.get("label") or value.get("type") or "").strip()
        content = str(value.get("block_content") or value.get("content") or value.get("text") or value.get("html") or "").strip()
        if label or content:
            block: dict[str, Any] = {
                "type": label or "block",
                "source_artifact": str(source_artifact),
            }
            if content:
                block["text_preview"] = content[:1000]
                block["text_char_count"] = len(content)
            bbox = normalize_jsonable(value.get("bbox") or value.get("box") or value.get("coordinate") or value.get("poly"))
            if bbox is not None:
                block["bbox"] = bbox
            confidence = float_or_none(value.get("confidence") or value.get("score"))
            if confidence is not None:
                block["confidence"] = confidence
            blocks.append(block)
        for child in value.values():
            blocks.extend(document_blocks_from_json_value(child, source_artifact))
    elif isinstance(value, list):
        for child in value:
            blocks.extend(document_blocks_from_json_value(child, source_artifact))
    return blocks


def table_candidates_from_raw(root: Path) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for path in sorted(root.rglob("*.json")):
        if path.name in GENERATED_SIDECAR_NAMES:
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        except Exception:
            continue
        tables.extend(table_candidates_from_json_value(data, path))
    for markdown, path in docx_table_candidates(root):
        tables.append({"label": "table", "markdown": markdown, "source_artifact": str(path)})
    return unique_table_candidates(tables)


def table_candidates_from_json_value(value: Any, source_artifact: Path) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    if isinstance(value, dict):
        label = str(value.get("block_label") or value.get("label") or value.get("type") or "").lower()
        content = str(value.get("block_content") or value.get("content") or value.get("html") or value.get("text") or "").strip()
        if "table" in label and content:
            table: dict[str, Any] = {"label": label or "table", "source_artifact": str(source_artifact)}
            if "<table" in content.lower():
                table["html"] = content
            else:
                table["markdown"] = normalize_table_content(content)
            bbox = normalize_jsonable(value.get("bbox") or value.get("box") or value.get("coordinate") or value.get("poly"))
            if bbox is not None:
                table["bbox"] = bbox
            confidence = float_or_none(value.get("confidence") or value.get("score"))
            if confidence is not None:
                table["confidence"] = confidence
            tables.append(table)
        for child in value.values():
            tables.extend(table_candidates_from_json_value(child, source_artifact))
    elif isinstance(value, list):
        for child in value:
            tables.extend(table_candidates_from_json_value(child, source_artifact))
    return tables


def docx_table_candidates(root: Path) -> list[tuple[str, Path]]:
    try:
        from docx import Document
    except Exception:
        return []
    tables: list[tuple[str, Path]] = []
    for path in sorted(root.rglob("*.docx")):
        try:
            document = Document(str(path))
        except Exception:
            continue
        for table in document.tables:
            rows = [[clean_cell_text(cell.text) for cell in row.cells] for row in table.rows]
            markdown = table_rows_to_markdown(rows)
            if markdown:
                tables.append((markdown, path))
    return tables


def unique_table_candidates(tables: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen = set()
    unique = []
    for table in tables:
        key = json.dumps(table, ensure_ascii=False, sort_keys=True)
        if key in seen:
            continue
        seen.add(key)
        unique.append(table)
    return unique


def normalize_jsonable(value: Any) -> Any | None:
    if value is None:
        return None
    if hasattr(value, "tolist"):
        value = value.tolist()
    if isinstance(value, dict):
        return {str(key): normalize_jsonable(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [normalize_jsonable(item) for item in value]
    if isinstance(value, (str, int, float, bool)):
        return value
    return str(value)


def float_or_none(value: Any) -> float | None:
    try:
        if value is None or value == "":
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def configure_local_cache(env: dict[str, str]) -> None:
    env.setdefault("HOME", str(TOOL_CACHE / "vlm-home"))
    env.setdefault("USERPROFILE", str(TOOL_CACHE / "vlm-home"))
    env.setdefault("XDG_CACHE_HOME", str(TOOL_CACHE / "vlm-cache"))
    env.setdefault("PADDLE_HOME", str(TOOL_CACHE / "paddle-cache"))
    env.setdefault("PADDLE_PDX_CACHE_HOME", str(TOOL_CACHE / "paddlex-cache"))
    env.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
    env.setdefault("HF_HUB_DISABLE_SYMLINKS", "1")
    env.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")
    for key in ("HOME", "USERPROFILE", "XDG_CACHE_HOME", "PADDLE_HOME", "PADDLE_PDX_CACHE_HOME"):
        Path(env[key]).mkdir(parents=True, exist_ok=True)


def pick_text_artifact(root: Path) -> Path | None:
    preferred = sorted(root.rglob("*.md")) + sorted(root.rglob("*.markdown")) + sorted(root.rglob("*.txt"))
    for path in preferred:
        try:
            if path.read_text(encoding="utf-8", errors="replace").strip():
                return path
        except Exception:
            continue
    return None


def write_normalized_markdown(candidate: Path, output: Path, raw_root: Path) -> None:
    markdown = candidate.read_text(encoding="utf-8", errors="replace").strip()
    tables = extract_actual_tables(raw_root)
    if tables:
        markdown = markdown.rstrip() + "\n\n## 真实表格 / Extracted Tables\n\n" + "\n\n".join(tables).strip()
    output.write_text(markdown.rstrip() + "\n", encoding="utf-8", newline="\n")


def extract_actual_tables(root: Path) -> list[str]:
    tables: list[str] = []
    tables.extend(extract_json_table_blocks(root))
    tables.extend(extract_docx_tables(root))
    return unique_markdown_blocks(tables)


def extract_json_table_blocks(root: Path) -> list[str]:
    tables: list[str] = []
    for path in sorted(root.rglob("*.json")):
        if path.name in GENERATED_SIDECAR_NAMES:
            continue
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
        except Exception:
            continue
        tables.extend(table_blocks_from_json_value(data))
    return tables


def table_blocks_from_json_value(value) -> list[str]:
    tables: list[str] = []
    if isinstance(value, dict):
        label = str(value.get("block_label") or value.get("label") or value.get("type") or "").lower()
        content = str(value.get("block_content") or value.get("content") or value.get("html") or "").strip()
        if "table" in label and content:
            tables.append(normalize_table_content(content))
        for child in value.values():
            tables.extend(table_blocks_from_json_value(child))
    elif isinstance(value, list):
        for child in value:
            tables.extend(table_blocks_from_json_value(child))
    return tables


def extract_docx_tables(root: Path) -> list[str]:
    try:
        from docx import Document
    except Exception:
        return []
    tables: list[str] = []
    for path in sorted(root.rglob("*.docx")):
        try:
            document = Document(str(path))
        except Exception:
            continue
        for table in document.tables:
            rows = [[clean_cell_text(cell.text) for cell in row.cells] for row in table.rows]
            markdown = table_rows_to_markdown(rows)
            if markdown:
                tables.append(markdown)
    return tables


def table_rows_to_markdown(rows: list[list[str]]) -> str:
    rows = [[cell for cell in row] for row in rows if any(cell.strip() for cell in row)]
    if not rows:
        return ""
    col_count = max(len(row) for row in rows)
    normalized = [row + [""] * (col_count - len(row)) for row in rows]
    if col_count < 2:
        return ""
    header = normalized[0]
    body = normalized[1:] or [[""] * col_count]
    lines = [
        "| " + " | ".join(markdown_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(col_count)) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(markdown_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def normalize_table_content(content: str) -> str:
    content = content.strip()
    if not content:
        return ""
    return content


def clean_cell_text(value: str) -> str:
    return " ".join(str(value).replace("\r", "\n").split())


def markdown_cell(value: str) -> str:
    return clean_cell_text(value).replace("|", r"\|")


def unique_markdown_blocks(blocks: list[str]) -> list[str]:
    seen = set()
    unique = []
    for block in blocks:
        normalized = "\n".join(line.rstrip() for line in str(block).strip().splitlines()).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(normalized)
    return unique


if __name__ == "__main__":
    raise SystemExit(main())
